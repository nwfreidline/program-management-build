"""Export logic for Career Tracker.

Handles exporting entries to Word, Excel, Markdown, and plain text formats.
"""
from pathlib import Path

from star_engine import format_entry_as_markdown, format_entry_as_text
from config import MONTH_NAMES


def _group_by_month(entries: list) -> dict:
    """Group entries by their date_completed (YYYY-MM) for organized export."""
    groups = {}
    for entry in entries:
        date_key = entry.get("date_completed", "Unknown")
        if date_key not in groups:
            groups[date_key] = []
        groups[date_key].append(entry)

    # Sort by date descending
    sorted_groups = dict(
        sorted(groups.items(), key=lambda x: x[0], reverse=True)
    )
    return sorted_groups


def _group_by_status(entries: list) -> dict:
    """Group entries into Completed and In Progress/Ongoing sections."""
    completed = []
    in_progress = []

    for entry in entries:
        status = entry.get("status", "").lower()
        if status in ("completed", "launched"):
            completed.append(entry)
        else:
            in_progress.append(entry)

    groups = {}
    if in_progress:
        groups["In Progress"] = in_progress
    if completed:
        groups["Completed"] = completed
    return groups


def _clean_entry_for_export(entry: dict) -> dict:
    """Clean an entry for export — strip placeholder text, ensure no [Add:] lines."""
    cleaned = dict(entry)

    # Strip placeholder fields
    for field in ("situation", "task", "result"):
        val = cleaned.get(field, "")
        if "[" in val and "Add:" in val:
            cleaned[field] = ""

    # Strip placeholder actions
    cleaned["actions"] = [
        a for a in cleaned.get("actions", [])
        if "[Add:" not in a
    ]

    return cleaned


def _format_date_heading(date_str: str) -> str:
    """Convert YYYY-MM to a readable heading like 'May 2026'."""
    try:
        parts = date_str.split("-")
        year = parts[0]
        month_num = int(parts[1]) if len(parts) > 1 else 0
        month_name = MONTH_NAMES[month_num] if 0 < month_num <= 12 else ""
        return f"{month_name} {year}".strip()
    except (ValueError, IndexError):
        return date_str


def export_to_markdown(entries: list, filepath: str):
    """Export all entries as a formatted markdown file grouped by status then month.

    Args:
        entries: List of entry dictionaries
        filepath: Output file path (.md)
    """
    lines = []
    lines.append("# Career Tracker — Accomplishments")
    lines.append("")

    completed_count = sum(1 for e in entries if e.get("status", "").lower() in ("completed", "launched"))
    in_progress_count = len(entries) - completed_count
    lines.append(f"*{len(entries)} entries ({completed_count} completed, {in_progress_count} in progress)*")
    lines.append("")

    status_groups = _group_by_status(entries)

    for status_label, status_entries in status_groups.items():
        lines.append(f"# {status_label}")
        lines.append("")

        month_groups = _group_by_month(status_entries)

        for date_key, group_entries in month_groups.items():
            heading = _format_date_heading(date_key)
            lines.append(f"## {heading}")
            lines.append("")

            for entry in group_entries:
                entry = _clean_entry_for_export(entry)
                lines.append(format_entry_as_markdown(entry))

    output = "\n".join(lines)
    Path(filepath).write_text(output, encoding="utf-8")


def export_to_text(entries: list, filepath: str):
    """Export all entries as a plain text file.

    Args:
        entries: List of entry dictionaries
        filepath: Output file path (.txt)
    """
    lines = []
    lines.append("CAREER TRACKER — ACCOMPLISHMENTS")
    lines.append(f"{len(entries)} entries")
    lines.append("=" * 60)
    lines.append("")

    groups = _group_by_month(entries)

    for date_key, group_entries in groups.items():
        heading = _format_date_heading(date_key)
        lines.append(heading.upper())
        lines.append("-" * 40)
        lines.append("")

        for entry in group_entries:
            lines.append(format_entry_as_text(entry))

    output = "\n".join(lines)
    Path(filepath).write_text(output, encoding="utf-8")


def export_to_word(entries: list, filepath: str):
    """Export entries to a styled Word document with STAR formatting.

    Groups entries by status (In Progress first, then Completed), then by month.
    Bold labels for STAR fields. Strips placeholder [Add:] text.

    Args:
        entries: List of entry dictionaries
        filepath: Output file path (.docx)
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx is required for Word export. "
            "Install it with: pip install python-docx"
        )

    doc = Document()

    # Title
    doc.add_heading("Career Tracker — Accomplishments", level=0)

    # Summary
    completed_count = sum(1 for e in entries if e.get("status", "").lower() in ("completed", "launched"))
    in_progress_count = len(entries) - completed_count
    subtitle = doc.add_paragraph(
        f"{len(entries)} entries ({completed_count} completed, {in_progress_count} in progress)"
    )
    subtitle.style.font.size = Pt(11)

    # Group by status, then by month within each group
    status_groups = _group_by_status(entries)

    for status_label, status_entries in status_groups.items():
        doc.add_heading(status_label, level=1)

        month_groups = _group_by_month(status_entries)

        for date_key, group_entries in month_groups.items():
            heading = _format_date_heading(date_key)
            doc.add_heading(heading, level=2)

            for entry in group_entries:
                # Clean placeholders before export
                entry = _clean_entry_for_export(entry)

                # Entry title
                doc.add_heading(entry["title"], level=3)

                # Status and date
                meta_para = doc.add_paragraph()
                run = meta_para.add_run("Status: ")
                run.bold = True
                meta_para.add_run(entry.get("status", "N/A"))
                meta_para.add_run("  |  ")
                run = meta_para.add_run("Date: ")
                run.bold = True
                meta_para.add_run(entry.get("date_completed", "N/A"))

                # Situation
                if entry.get("situation"):
                    sit_para = doc.add_paragraph()
                    run = sit_para.add_run("Situation: ")
                    run.bold = True
                    sit_para.add_run(entry["situation"])

                # Task
                if entry.get("task"):
                    task_para = doc.add_paragraph()
                    run = task_para.add_run("Task: ")
                    run.bold = True
                    task_para.add_run(entry["task"])

                # Actions
                actions = entry.get("actions", [])
                if actions:
                    actions_para = doc.add_paragraph()
                    run = actions_para.add_run("Actions:")
                    run.bold = True
                    for action in actions:
                        doc.add_paragraph(action, style="List Bullet")

                # Result
                if entry.get("result"):
                    result_para = doc.add_paragraph()
                    run = result_para.add_run("Result: ")
                    run.bold = True
                    result_para.add_run(entry["result"])

                # Separator
                doc.add_paragraph()

    doc.save(filepath)


def export_to_excel(entries: list, filepath: str):
    """Export entries to an Excel spreadsheet.

    One row per entry with columns for each STAR field.

    Args:
        entries: List of entry dictionaries
        filepath: Output file path (.xlsx)

    Raises:
        ImportError: If openpyxl is not installed
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
    except ImportError:
        raise ImportError(
            "openpyxl is required for Excel export. "
            "Install it with: pip install openpyxl"
        )

    wb = Workbook()
    ws = wb.active
    ws.title = "Career Entries"

    # Headers
    headers = ["Title", "Status", "Date", "Situation", "Task", "Actions", "Result"]
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_text_font = Font(bold=True, size=11, color="FFFFFF")

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_text_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Data rows
    for row_idx, entry in enumerate(entries, 2):
        ws.cell(row=row_idx, column=1, value=entry.get("title", ""))
        ws.cell(row=row_idx, column=2, value=entry.get("status", ""))
        ws.cell(row=row_idx, column=3, value=entry.get("date_completed", ""))
        ws.cell(row=row_idx, column=4, value=entry.get("situation", ""))
        ws.cell(row=row_idx, column=5, value=entry.get("task", ""))

        # Join actions with newlines
        actions = entry.get("actions", [])
        ws.cell(row=row_idx, column=6, value="\n".join(actions))

        ws.cell(row=row_idx, column=7, value=entry.get("result", ""))

        # Wrap text for long fields
        for col in range(4, 8):
            ws.cell(row=row_idx, column=col).alignment = Alignment(
                wrap_text=True, vertical="top"
            )

    # Auto-width columns (approximate)
    column_widths = [30, 15, 12, 40, 40, 50, 40]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[chr(64 + i)].width = width

    wb.save(filepath)
