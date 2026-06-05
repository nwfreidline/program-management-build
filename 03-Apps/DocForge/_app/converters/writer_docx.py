"""Write internal document model to Word (.docx), with optional template support."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _set_cell_border(cell, color="BFBFBF", size="4"):
    """Add light borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    """Set compact cell margins (in twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _shade_cell(cell, color="F2F2F2"):
    """Apply background shading to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color)
    tcPr.append(shading)


def _format_cell_text(cell, text, font_name="Calibri", font_size=10,
                      bold=False, color=None):
    """Set cell text with specific font formatting."""
    # Clear all existing paragraphs/runs
    for para in cell.paragraphs:
        para.clear()
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Vertically center cell content
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _clear_body(doc):
    """Remove all paragraphs and tables from document body."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)


def _resolve_style(doc, preferred, fallback="Normal"):
    """Try to use a style name; fall back if not found."""
    try:
        doc.styles[preferred]
        return preferred
    except KeyError:
        return fallback


def _ensure_numbering(doc):
    """Ensure the document has bullet and number list definitions.
    
    Creates numbering definitions if they don't exist:
    - numId 1: bullet list (•)
    - numId 2: numbered list (1. 2. 3.)
    """
    # Access or create the numbering part
    numbering_part = doc.part.numbering_part
    numbering = numbering_part.numbering_definitions._numbering

    # Check if numId 1 and 2 already exist
    existing_nums = numbering.findall(qn("w:num"))
    existing_ids = {el.get(qn("w:numId")) for el in existing_nums}

    if "1" not in existing_ids:
        # Create abstract numbering for bullets
        abstract_bullet = OxmlElement("w:abstractNum")
        abstract_bullet.set(qn("w:abstractNumId"), "0")
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "bullet")
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "\u2022")
        lvl.append(lvlText)
        lvlJc = OxmlElement("w:lvlJc")
        lvlJc.set(qn("w:val"), "left")
        lvl.append(lvlJc)
        abstract_bullet.append(lvl)
        numbering.insert(0, abstract_bullet)

        num_bullet = OxmlElement("w:num")
        num_bullet.set(qn("w:numId"), "1")
        abstractNumId = OxmlElement("w:abstractNumId")
        abstractNumId.set(qn("w:val"), "0")
        num_bullet.append(abstractNumId)
        numbering.append(num_bullet)

    if "2" not in existing_ids:
        # Create abstract numbering for numbers
        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), "1")
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "decimal")
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "%1.")
        lvl.append(lvlText)
        lvlJc = OxmlElement("w:lvlJc")
        lvlJc.set(qn("w:val"), "left")
        lvl.append(lvlJc)
        abstract_num.append(lvl)
        numbering.insert(0, abstract_num)

        num_num = OxmlElement("w:num")
        num_num.set(qn("w:numId"), "2")
        abstractNumId = OxmlElement("w:abstractNumId")
        abstractNumId.set(qn("w:val"), "1")
        num_num.append(abstractNumId)
        numbering.append(num_num)


def _apply_bullet_format(paragraph, level=0):
    """Apply bullet list formatting to a paragraph using Word numbering XML."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numPr.append(ilvl)
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(numId)
    pPr.append(numPr)
    paragraph.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)


def _apply_number_format(paragraph, level=0, number=1):
    """Apply numbered list formatting to a paragraph using Word numbering XML."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numPr.append(ilvl)
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "2")
    numPr.append(numId)
    pPr.append(numPr)
    paragraph.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)


def _apply_default_styles(doc):
    """Apply default formatting to a new document (no template).
    
    Standardizes to:
    - Calibri 10pt body text, black, single-spaced, 6pt after
    - Headings: Calibri bold, black (no color), scaled sizes
    - H1: 14pt, H2: 12pt, H3: 10pt bold, H4: 10pt bold italic
    - List Paragraph: 10pt, 0.25" indent
    - Page margins: 1" all sides (standard)
    """
    # Page setup
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Normal / body
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.bold = False
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.0

    # Headings — Calibri, bold, black, no color
    heading_sizes = {1: 14, 2: 12, 3: 10, 4: 10}
    for level, size in heading_sizes.items():
        style_name = f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = (level == 4)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    # List Paragraph
    try:
        list_style = doc.styles["List Paragraph"]
        list_style.font.name = "Calibri"
        list_style.font.size = Pt(10)
        list_style.font.color.rgb = RGBColor(0, 0, 0)
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(12)
        list_style.paragraph_format.left_indent = Inches(0.25)
        list_style.paragraph_format.line_spacing = 1.0
    except KeyError:
        pass


def write_docx(doc_model, output_path, template_path=None):
    """Convert document model to a .docx file.
    
    Args:
        doc_model: The internal document dict
        output_path: Where to save the .docx
        template_path: Optional .docx template to use as style source
    """
    if template_path:
        doc = Document(template_path)
        _clear_body(doc)
        # Update header if template has one
        for section in doc.sections:
            header = section.header
            if header.paragraphs and header.paragraphs[0].text:
                header_para = header.paragraphs[0]
                header_para.clear()
                run = header_para.add_run(f"DocForge \u2014 {doc_model['title']}")
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
    else:
        doc = Document()
        _apply_default_styles(doc)

    # Ensure bullet and number list definitions exist
    _ensure_numbering(doc)

    # Detect available styles
    heading_styles = {}
    for level in range(1, 5):
        heading_styles[level] = _resolve_style(doc, f"Heading {level}", "Normal")

    bullet_style = _resolve_style(doc, "List Bullet",
                       _resolve_style(doc, "List Paragraph", "Normal"))
    number_style = _resolve_style(doc, "List Number",
                       _resolve_style(doc, "List Paragraph", "Normal"))
    body_style = _resolve_style(doc, "Normal")

    # Track numbering for numbered lists
    _num_counter = [0]
    _prev_was_numbered = [False]

    for block in doc_model["blocks"]:
        btype = block["type"]

        # Reset number counter when leaving a numbered list
        if btype != "numbered" and _prev_was_numbered[0]:
            _num_counter[0] = 0
            _prev_was_numbered[0] = False

        if btype == "heading":
            level = block.get("level", 1)
            style = heading_styles.get(level, heading_styles[1])
            doc.add_paragraph(block["text"], style=style)

        elif btype == "paragraph":
            doc.add_paragraph(block["text"], style=body_style)

        elif btype == "bullet":
            p = doc.add_paragraph(style=bullet_style)
            indent_level = block.get("level", 0)
            # Apply bullet formatting via numbering XML
            _apply_bullet_format(p, indent_level)
            run = p.add_run(block["text"])
            run.font.name = "Calibri"
            run.font.size = Pt(10)

        elif btype == "numbered":
            _num_counter[0] += 1
            _prev_was_numbered[0] = True
            p = doc.add_paragraph(style=number_style)
            indent_level = block.get("level", 0)
            # Apply number formatting via numbering XML
            _apply_number_format(p, indent_level, _num_counter[0])
            run = p.add_run(block["text"])
            run.font.name = "Calibri"
            run.font.size = Pt(10)

        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for r_idx, row in enumerate(rows):
                for c_idx in range(num_cols):
                    cell = table.rows[r_idx].cells[c_idx]
                    val = row[c_idx] if c_idx < len(row) else ""

                    _set_cell_border(cell)
                    _set_cell_margins(cell)

                    if r_idx == 0:
                        # Header row: bold, light gray background
                        _shade_cell(cell, "F2F2F2")
                        _format_cell_text(cell, val, bold=True, font_size=10)
                    else:
                        _format_cell_text(cell, val, bold=False, font_size=10)

            # Small spacer after table
            spacer = doc.add_paragraph("", style=body_style)
            spacer.paragraph_format.space_before = Pt(2)
            spacer.paragraph_format.space_after = Pt(6)

        elif btype == "code":
            # Monospace code block with light background shading
            p = doc.add_paragraph(style=body_style)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)

    doc.save(output_path)
