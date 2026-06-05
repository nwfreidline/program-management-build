"""
Growth Timeline converter — parses growth_timeline markdown and writes
a polished .docx using the Growth Timeline template styles.

Handles the actual markdown structure:
- # title → document title (centered, blue)
- > blockquote after title → subtitle
- ## BRAINSTORM & ONGOING / ## COMPLETED → summary sections with bullet lists
- #### entries (or ## entries with STAR content) → project entries
- **Situation/Task/Action/Result:** → STAR labels with content
- Indented bullets under Action/Result → bulleted lists
- > Supporting Evidence blocks → gray italic evidence callout
- Skips: "How to Use", "New Item Template", "Summary Statistics"
"""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Formatting helpers ───────────────────────────────────────────────────────

DARK_BLUE = RGBColor(0, 51, 102)


def _clear_body(doc):
    """Remove all paragraphs and tables from document body."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)


def _ensure_numbering(doc):
    """Set up small-dot bullet list definition (numId 1)."""
    numbering_part = doc.part.numbering_part
    numbering = numbering_part.numbering_definitions._numbering

    for existing in numbering.findall(qn("w:num")):
        numbering.remove(existing)
    for existing in numbering.findall(qn("w:abstractNum")):
        numbering.remove(existing)

    abstract_bullet = OxmlElement("w:abstractNum")
    abstract_bullet.set(qn("w:abstractNumId"), "0")

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "bullet")
    lvl.append(numFmt)

    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), "\u2219")
    lvl.append(lvlText)

    lvlJc = OxmlElement("w:lvlJc")
    lvlJc.set(qn("w:val"), "left")
    lvl.append(lvlJc)

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:hint"), "default")
    rPr.append(rFonts)
    lvl.append(rPr)

    pPr_lvl = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    pPr_lvl.append(ind)
    lvl.append(pPr_lvl)

    abstract_bullet.append(lvl)
    numbering.append(abstract_bullet)

    num_bullet = OxmlElement("w:num")
    num_bullet.set(qn("w:numId"), "1")
    abstractNumId = OxmlElement("w:abstractNumId")
    abstractNumId.set(qn("w:val"), "0")
    num_bullet.append(abstractNumId)
    numbering.append(num_bullet)


def _add_bottom_border(paragraph, color="003366", size="8"):
    """Add a bottom border accent to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bullet(doc, text, level=0):
    """Add a bullet point with small-dot formatting."""
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Inches(0.5 + (level * 0.25))
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numPr.append(ilvl)
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(numId)
    pPr.append(numPr)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def _add_star_label(doc, label, text):
    """Add a STAR-format line: bold blue label + normal text."""
    p = doc.add_paragraph(style="GT STAR")
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.name = "Calibri"
    run_label.font.color.rgb = DARK_BLUE
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)
    run_text.font.name = "Calibri"


def _add_star_label_only(doc, label):
    """Add just the STAR label without trailing text."""
    p = doc.add_paragraph(style="GT STAR")
    run = p.add_run(f"{label}:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_BLUE


def _add_evidence(doc, lines):
    """Add a supporting evidence block."""
    # "Supporting Evidence:" label aligned with STAR labels (same indent)
    p = doc.add_paragraph(style="GT Evidence")
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run("Supporting Evidence:")
    run.bold = True
    # Sub-lines stay at their current indent
    for line in lines:
        if line.strip():
            p = doc.add_paragraph(style="GT Evidence")
            p.add_run(line)


# ─── Markdown stripping ──────────────────────────────────────────────────────

def _strip_md(text):
    """Remove inline markdown formatting."""
    text = re.sub(r"\*{3}(.+?)\*{3}", r"\1", text)
    text = re.sub(r"_{3}(.+?)_{3}", r"\1", text)
    text = re.sub(r"\*{2}(.+?)\*{2}", r"\1", text)
    text = re.sub(r"_{2}(.+?)_{2}", r"\1", text)
    text = re.sub(r"(?<!\w)\*(.+?)\*(?!\w)", r"\1", text)
    text = re.sub(r"(?<!\w)_(.+?)_(?!\w)", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    return text.strip()


# ─── Parser ──────────────────────────────────────────────────────────────────

def _parse_growth_timeline(filepath):
    """Parse a growth timeline markdown into structured data.

    Returns:
        dict with:
            title: str
            subtitle: str
            sections: list of dicts, each with:
                heading: str
                type: 'summary' | 'entries'
                items: list (strings for summary, entry dicts for entries)

        Entry dict:
            name: str
            status: str
            situation: str
            task: str
            action: list of str
            result: list of str
            evidence: list of str
    """
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()

    data = {"title": "Growth Timeline", "subtitle": "", "sections": []}

    i = 0
    total = len(lines)

    # ── Title ──
    while i < total:
        line = lines[i].rstrip()
        if line.startswith("# ") and not line.startswith("## "):
            data["title"] = _strip_md(line[2:])
            i += 1
            break
        i += 1

    # ── Subtitle (> blockquote) ──
    while i < total and not lines[i].strip():
        i += 1
    subtitle_parts = []
    while i < total and lines[i].strip().startswith(">"):
        subtitle_parts.append(lines[i].strip().lstrip("> ").strip())
        i += 1
    if subtitle_parts:
        data["subtitle"] = " ".join(subtitle_parts)

    # ── Main parsing loop ──
    # Track state
    current_section = None  # current section dict
    current_entry = None    # current entry dict
    star_key = None         # which STAR field we're collecting for
    in_evidence = False
    evidence_buf = []
    skip_section = False    # skip "How to Use", "Summary Statistics", etc.

    SKIP_SECTIONS = {"how to use", "summary statistics", "new item template",
                     "project/task/achievement"}
    SUMMARY_KEYWORDS = {"completed", "brainstorm", "ongoing", "in progress"}
    ENTRIES_KEYWORDS = {"project history", "may 2026", "april 2026", "march 2026",
                        "february 2026", "january 2026", "june 2026", "july 2026",
                        "august 2026", "september 2026", "october 2026",
                        "november 2026", "december 2026"}

    def _flush_entry():
        """Save current entry to current section."""
        nonlocal current_entry, in_evidence, evidence_buf, star_key
        if current_entry is None:
            return
        if in_evidence:
            current_entry["evidence"] = evidence_buf
            in_evidence = False
            evidence_buf = []
        star_key = None
        # Only add if it has a name
        if current_entry.get("name"):
            if current_section is None:
                _ensure_entries_section()
            current_section["items"].append(current_entry)
        current_entry = None

    def _flush_section():
        """Save current section to data."""
        nonlocal current_section
        _flush_entry()
        if current_section and current_section.get("items"):
            data["sections"].append(current_section)
        current_section = None

    def _ensure_entries_section():
        """Create an implicit Project History section if none exists."""
        nonlocal current_section
        if current_section is None or current_section["type"] != "entries":
            _flush_section()
            current_section = {
                "heading": "Project History",
                "type": "entries",
                "items": [],
            }

    def _is_summary_heading(heading):
        """Check if a heading is a summary section (COMPLETED, BRAINSTORM, etc.)."""
        h_lower = heading.lower()
        return any(kw in h_lower for kw in SUMMARY_KEYWORDS)

    def _is_entries_heading(heading):
        """Check if a heading is an entries section header (Project History, month names)."""
        h_lower = heading.lower()
        return any(kw in h_lower for kw in ENTRIES_KEYWORDS)

    def _should_skip(heading):
        """Check if a section should be skipped."""
        h_lower = heading.lower()
        return any(kw in h_lower for kw in SKIP_SECTIONS)

    while i < total:
        line = lines[i].rstrip()

        # ── Blank line ──
        if not line.strip():
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r"^[-*_]{3,}\s*$", line):
            i += 1
            continue

        # ── ## Section heading ──
        if line.startswith("## ") and not line.startswith("### "):
            heading = _strip_md(line[3:])

            if _should_skip(heading):
                # Skip this entire section until next ## or ####
                _flush_entry()
                skip_section = True
                i += 1
                continue

            skip_section = False

            if _is_summary_heading(heading):
                _flush_section()
                current_section = {
                    "heading": heading,
                    "type": "summary",
                    "items": [],
                }
                star_key = None
                i += 1
                continue
            elif _is_entries_heading(heading):
                # This is an explicit entries section header (e.g., "Project History")
                _flush_section()
                current_section = {
                    "heading": heading,
                    "type": "entries",
                    "items": [],
                }
                star_key = None
                i += 1
                continue
            else:
                # Could be an entry disguised as ## (like "CCS - Program...")
                # Check if next lines have STAR content
                # For now, treat as an entry under Project History
                _flush_entry()
                _ensure_entries_section()
                entry_name = heading
                # Strip known prefixes like "Narrative - "
                _STRIP_PREFIXES = {"narrative", "quick suite"}
                for sep in [" - ", " — ", " – "]:
                    if sep in entry_name:
                        parts = entry_name.split(sep, 1)
                        if parts[0].strip().lower() in _STRIP_PREFIXES:
                            entry_name = parts[1]
                        break
                current_entry = {
                    "name": entry_name,
                    "status": "",
                    "situation": "",
                    "task": "",
                    "action": [],
                    "result": [],
                    "evidence": [],
                }
                star_key = None
                i += 1
                continue

        # ── Skip mode ──
        if skip_section:
            # Check if we hit a new section that should end skip mode
            if line.startswith("## ") or line.startswith("#### "):
                skip_section = False
                # Don't increment — re-process this line
                continue
            i += 1
            continue

        # ── #### Entry heading ──
        if line.startswith("#### "):
            heading = _strip_md(line[5:])

            # Skip "New Item Template"
            if "new item template" in heading.lower():
                skip_section = True
                _flush_entry()
                i += 1
                continue

            # If we were skipping, stop now
            skip_section = False

            _flush_entry()
            _ensure_entries_section()

            entry_name = heading
            # Strip known prefixes like "Narrative - "
            _STRIP_PREFIXES = {"narrative", "quick suite"}
            for sep in [" - ", " — ", " – "]:
                if sep in entry_name:
                    parts = entry_name.split(sep, 1)
                    if parts[0].strip().lower() in _STRIP_PREFIXES:
                        entry_name = parts[1]
                    break

            current_entry = {
                "name": entry_name,
                "status": "",
                "situation": "",
                "task": "",
                "action": [],
                "result": [],
                "evidence": [],
            }
            star_key = None
            i += 1
            continue

        # ── Status line (*Status Text*) ──
        if current_entry and re.match(r"^\*[^*]+\*\s*$", line):
            current_entry["status"] = line.strip().strip("*")
            i += 1
            continue

        # ── Blockquote (evidence or other) ──
        if line.strip().startswith(">"):
            content = line.strip()
            # Remove leading > and optional space
            content = re.sub(r"^>\s*", "", content)

            if "supporting evidence" in content.lower():
                in_evidence = True
                i += 1
                continue
            elif in_evidence:
                # Collect evidence lines
                cleaned = _strip_md(content.lstrip("- ").strip())
                if cleaned:
                    evidence_buf.append(cleaned)
                i += 1
                continue
            else:
                # Other blockquote (attaboy, notes, etc.) — skip for now
                i += 1
                continue

        # ── End of evidence block (non-blockquote line after evidence) ──
        if in_evidence and not line.strip().startswith(">"):
            if current_entry:
                current_entry["evidence"] = evidence_buf
            in_evidence = False
            evidence_buf = []
            # Don't increment — re-process this line

        # ── STAR label: - **Situation:** / - **Task:** / - **Action:** / - **Result:** ──
        star_match = re.match(
            r"^[-\s]*\*\*(Situation|Task|Action|Result)\s*:\*\*\s*(.*)", line
        )
        if star_match and current_entry:
            label = star_match.group(1).lower()
            content = _strip_md(star_match.group(2))
            star_key = label

            if label in ("situation", "task"):
                current_entry[label] = content
            elif label == "action":
                if content:
                    current_entry["action"].append(content)
            elif label == "result":
                if content:
                    current_entry["result"].append(content)
            i += 1
            continue

        # ── Bullet items ──
        bullet_match = re.match(r"^(\s*)[*\-+]\s+(.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = _strip_md(bullet_match.group(2))

            if current_entry and star_key == "action":
                current_entry["action"].append(text)
            elif current_entry and star_key == "result":
                current_entry["result"].append(text)
            elif current_entry and star_key in ("situation", "task"):
                # Continuation bullet under situation/task (rare)
                current_entry[star_key] += f" {text}"
            elif current_section and current_section["type"] == "summary":
                current_section["items"].append(text)
            elif current_entry:
                # Generic bullet — add to action by default
                current_entry["action"].append(text)
            i += 1
            continue

        # ── Table (skip) ──
        if "|" in line and i + 1 < total and re.match(r"^\s*\|[-:\s|]+\|\s*$", lines[i + 1].rstrip()):
            # Skip table rows
            i += 1  # header
            i += 1  # separator
            while i < total and "|" in lines[i]:
                i += 1
            continue

        # ── Continuation text ──
        if line.strip() and current_entry and star_key:
            content = _strip_md(line)
            if star_key in ("situation", "task"):
                if current_entry[star_key]:
                    current_entry[star_key] += f" {content}"
                else:
                    current_entry[star_key] = content
            elif star_key == "action":
                current_entry["action"].append(content)
            elif star_key == "result":
                current_entry["result"].append(content)

        i += 1

    # Finalize
    _flush_section()

    return data


# ─── Writer ──────────────────────────────────────────────────────────────────

def write_growth_timeline(input_path, output_path, template_path):
    """Convert a growth timeline markdown to a styled .docx.

    Args:
        input_path: Path to the source .md file
        output_path: Where to save the .docx
        template_path: Path to Growth_Timeline_Template.docx
    """
    data = _parse_growth_timeline(input_path)

    doc = Document(template_path)
    _clear_body(doc)
    _ensure_numbering(doc)

    # === TITLE ===
    title_para = doc.add_paragraph(style="GT Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run(data["title"])

    # === SUBTITLE ===
    if data["subtitle"]:
        subtitle_para = doc.add_paragraph(style="GT Subtitle")
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.add_run(data["subtitle"])

    # === SECTIONS ===
    for section in data["sections"]:
        # Section header with bottom border
        h1 = doc.add_paragraph(section["heading"], style="Heading 1")
        _add_bottom_border(h1)

        if section["type"] == "summary":
            for item in section["items"]:
                if isinstance(item, str):
                    _add_bullet(doc, item)
        else:
            # Entries section
            for idx, entry in enumerate(section["items"]):
                if not isinstance(entry, dict):
                    continue

                # Add extra spacing between entries (not before the first one)
                if idx > 0:
                    spacer = doc.add_paragraph(style="Normal")
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)

                # Entry title
                entry_para = doc.add_paragraph(style="Heading 2")
                entry_para.add_run(entry["name"])

                # Status
                if entry.get("status"):
                    status_para = doc.add_paragraph(style="GT Status")
                    status_para.add_run(entry["status"])

                # Situation
                if entry.get("situation"):
                    _add_star_label(doc, "Situation", entry["situation"])

                # Task
                if entry.get("task"):
                    _add_star_label(doc, "Task", entry["task"])

                # Action
                if entry.get("action"):
                    _add_star_label_only(doc, "Action")
                    for item in entry["action"]:
                        _add_bullet(doc, item)

                # Result
                if entry.get("result"):
                    _add_star_label_only(doc, "Result")
                    for item in entry["result"]:
                        _add_bullet(doc, item)

                # Evidence
                if entry.get("evidence"):
                    _add_evidence(doc, entry["evidence"])

    doc.save(output_path)
