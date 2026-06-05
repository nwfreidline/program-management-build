"""
Preview script — generates a sample Growth Timeline .docx using the template.

This creates a mock output showing how growth_timeline_2026.md would render
with the Growth Timeline template applied. Use this to review formatting
before integrating into DocForge's converter pipeline.

Run: python preview_growth_timeline.py
Output: Growth_Timeline_Preview.docx (in this directory)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_bottom_border(paragraph, color="003366", size="8"):
    """Add a bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_star_paragraph(doc, label, text):
    """Add a STAR-format line: bold label (dark blue) + normal text."""
    p = doc.add_paragraph(style="GT STAR")
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.name = "Calibri"
    run_label.font.color.rgb = RGBColor(0, 51, 102)
    run_text = p.add_run(text)
    run_text.font.size = Pt(11)
    run_text.font.name = "Calibri"


def add_bullet(doc, text, level=0):
    """Add a bullet point with proper Word bullet formatting (small dot)."""
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Inches(0.5 + (level * 0.25))
    # Apply bullet numbering via XML
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numPr.append(ilvl)
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(numId)
    pPr.append(numPr)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def add_evidence_block(doc, lines):
    """Add a supporting evidence blockquote."""
    p = doc.add_paragraph(style="GT Evidence")
    run = p.add_run("Supporting Evidence:")
    run.bold = True
    for line in lines:
        p = doc.add_paragraph(style="GT Evidence")
        p.add_run(line)


def _ensure_numbering(doc):
    """Ensure the document has a small-dot bullet list definition (numId 1)."""
    numbering_part = doc.part.numbering_part
    numbering = numbering_part.numbering_definitions._numbering

    # Remove any existing num definitions to start clean
    for existing in numbering.findall(qn("w:num")):
        numbering.remove(existing)
    for existing in numbering.findall(qn("w:abstractNum")):
        numbering.remove(existing)

    # Create abstract numbering for small dot bullets
    abstract_bullet = OxmlElement("w:abstractNum")
    abstract_bullet.set(qn("w:abstractNumId"), "0")
    
    # Level 0 — small dot
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    
    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "bullet")
    lvl.append(numFmt)
    
    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), "\u2219")  # bullet operator — small filled dot
    lvl.append(lvlText)
    
    lvlJc = OxmlElement("w:lvlJc")
    lvlJc.set(qn("w:val"), "left")
    lvl.append(lvlJc)
    
    # Font for the bullet character
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:hint"), "default")
    rPr.append(rFonts)
    lvl.append(rPr)
    
    # Paragraph indent
    pPr_lvl = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")    # 0.5 inch total
    ind.set(qn("w:hanging"), "360") # bullet hangs 0.25 inch
    pPr_lvl.append(ind)
    lvl.append(pPr_lvl)
    
    abstract_bullet.append(lvl)
    numbering.append(abstract_bullet)

    # Link numId 1 to abstractNumId 0
    num_bullet = OxmlElement("w:num")
    num_bullet.set(qn("w:numId"), "1")
    abstractNumId = OxmlElement("w:abstractNumId")
    abstractNumId.set(qn("w:val"), "0")
    num_bullet.append(abstractNumId)
    numbering.append(num_bullet)


def build_preview():
    # Load the template
    template_path = Path(__file__).parent / "docx" / "Growth_Timeline_Template.docx"
    doc = Document(str(template_path))

    # Ensure bullet numbering definitions exist
    _ensure_numbering(doc)

    # Clear placeholder content
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)

    # === TITLE SECTION ===
    title = doc.add_paragraph(style="GT Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Growth Timeline \u2014 2026")

    subtitle = doc.add_paragraph(style="GT Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "A historical record of completed work across key projects, formatted in "
        "STAR format for use in performance reviews, promotion discussions, and "
        "compensation conversations. Entries are in reverse-chronological order."
    )

    # === COMPLETED SUMMARY ===
    doc.add_paragraph("Completed Projects", style="Heading 1")
    add_bottom_border(doc.paragraphs[-1])

    completed_items = [
        "AMCOP & BACOP PM Ingestion",
        "Fire Department False Alarms \u2014 Audit, Leadership Reporting & Vendor Financial Recovery",
        "Maintenance Reports \u2014 Program Launch & Ongoing Operations",
        "PM On-boarding Playbook",
        "Maintenance Dashboard \u2014 Comprehensive EAM & MCM",
        "PM Training Powerpoint overhaul",
        "PM Maintenance Program Training",
        "CCS Program Management",
        "NSL Program - PDX Tracking & Wiki Feedback",
        "PO & Invoice Dashboard",
        "Cost Savings Review Agent",
        "Invoice Review Agent",
    ]
    for item in completed_items:
        add_bullet(doc, item)

    # === IN PROGRESS / BRAINSTORM ===
    doc.add_paragraph("In Progress & Brainstorm", style="Heading 1")
    add_bottom_border(doc.paragraphs[-1])

    brainstorm_items = [
        "2027 Maintenance Calendar Overhaul",
        "PDC Ingestion",
        "RSL & IPS New Site Launch Procurement Ingestion",
        "Maintenance Report Platform",
        "Vendor Feedback Tool",
        "External Optics Access Platform",
    ]
    for item in brainstorm_items:
        add_bullet(doc, item)

    # === PROJECT HISTORY SECTION ===
    doc.add_paragraph("Project History", style="Heading 1")
    add_bottom_border(doc.paragraphs[-1])

    # --- Entry 1: Maintenance Report Desktop Application ---
    entry_title = doc.add_paragraph(style="Heading 2")
    entry_title.add_run("Maintenance Report Desktop Application")

    status = doc.add_paragraph(style="GT Status")
    status.add_run("Complete")

    add_star_paragraph(
        doc, "Situation",
        "The PM team\u2019s maintenance reporting process relied on manually populated "
        "Excel spreadsheets with complex macros, requiring each team member to navigate "
        "multi-step workflows for data entry, formatting, and file management. The process "
        "was error-prone, time-consuming, and inconsistent across team members."
    )

    add_star_paragraph(
        doc, "Task",
        "Design and build a desktop application that automates the maintenance report "
        "workflow \u2014 simplifying data entry, standardizing output formatting, and "
        "deploying the tool to all area-specific PDX folders for team-wide adoption."
    )

    # Action as bullet list
    p = doc.add_paragraph(style="GT STAR")
    run = p.add_run("Action:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)

    actions = [
        "Compiled all documents for current reporting process from PM Team folders",
        "Analyzed existing report templates (CCS, UPS, and general maintenance tracking) to understand data requirements",
        "Built Maintenance Report desktop app incorporating all required fields and automated formatting",
        "Tested desktop app with team members \u2014 gathered feedback on usability and edge cases",
        "Added copy of desktop app to each area-specific PDX folder",
        "Updated instructions on all master maintenance report spreadsheets",
        "Rolled out to team for use at their discretion",
    ]
    for action in actions:
        add_bullet(doc, action)

    add_star_paragraph(
        doc, "Result",
        "Desktop application built, tested, and deployed across all area-specific PDX "
        "maintenance report folders. Updated instructions distributed on all master report "
        "spreadsheets. Tool available to all team members, reducing manual data entry and "
        "formatting errors."
    )

    add_evidence_block(doc, [
        "\U0001f4c1 Folder: TPM-Support > Maintenance > Maintenance Reports > 0 Templates > Test Workspace"
    ])

    # --- Entry 2: PDX PM Wiki Overhaul ---
    entry_title = doc.add_paragraph(style="Heading 2")
    entry_title.add_run("PDX PM Wiki Overhaul")

    status = doc.add_paragraph(style="GT Status")
    status.add_run("Complete")

    add_star_paragraph(
        doc, "Situation",
        "The PDX Program Manager team wiki had grown organically without consistent "
        "formatting, theming, or information architecture. Pages used inconsistent layouts, "
        "outdated contact information, broken links, and lacked a cohesive visual identity."
    )

    add_star_paragraph(
        doc, "Task",
        "Perform a comprehensive overhaul of all PDX PM wiki pages \u2014 standardizing "
        "formatting and theme, updating all content to reflect current data, creating "
        "missing pages, and producing a reusable styling guide."
    )

    p = doc.add_paragraph(style="GT STAR")
    run = p.add_run("Action:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)

    actions = [
        "Created comprehensive wiki formatting steering document (15-section PDX Wiki Styling Guide)",
        "Created manual editing quick reference card for team members",
        "Standardized formatting and dark theme across all 15+ wiki pages",
        "Built reusable transclude templates for navigation bar and footer",
        "Updated all pages with current data (contact info, vendor info, links)",
        "Built the Resources > Training wiki page from scratch",
        "Documented all XWiki syntax patterns and workarounds for team reference",
    ]
    for action in actions:
        add_bullet(doc, action)

    add_star_paragraph(
        doc, "Result",
        "All 15+ PDX PM wiki pages reformatted with consistent dark theme and standardized "
        "layout. Comprehensive styling guide produced (15 sections). Training wiki page "
        "built from scratch. All contact information, vendor data, and links verified and updated."
    )

    add_evidence_block(doc, [
        "\U0001f4c1 Folder: TPM Kiro Projects > PM Wiki Overhaul"
    ])

    # --- Entry 3: Cost Savings Review Agent (abbreviated) ---
    entry_title = doc.add_paragraph(style="Heading 2")
    entry_title.add_run("Cost Savings Review Agent")

    status = doc.add_paragraph(style="GT Status")
    status.add_run("Complete")

    add_star_paragraph(
        doc, "Situation",
        "The PDX TPM team managed cost savings submissions from AMER DCEO teams via SIM "
        "tickets, requiring manual review against the DCC Cost Savings Guidance policy. "
        "Reviews were time-consuming, inconsistent, and lacked a standardized framework."
    )

    add_star_paragraph(
        doc, "Task",
        "Design and build an AI-powered cost savings review agent that performs structured "
        "first-pass evaluations, produces ready-to-use SIM comments, and generates "
        "Salesforce entry blocks for approved initiatives."
    )

    p = doc.add_paragraph(style="GT STAR")
    run = p.add_run("Action:")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)

    actions = [
        "Built comprehensive 18-document knowledge base from policy documentation",
        "Developed structured decision flow and steering guide",
        "Iterated through 5 agent persona versions with team feedback",
        "Configured agent to evaluate against 7 criteria with 4-tier decision framework",
        "Designed SIM auto-resolve rules (14-day window) for abandoned tickets",
        "Tested and launched for production use across 6 regions (PDX, SFO, QRO, GRU, YUL, YYC)",
    ]
    for action in actions:
        add_bullet(doc, action)

    add_star_paragraph(
        doc, "Result",
        "Fully operational AI review agent processing cost savings submissions across 6 "
        "regions. 18-document knowledge base built from scratch. 14-day auto-resolve "
        "automation eliminates abandoned ticket backlog. Standardized 4-tier decision "
        "framework with confidence scoring."
    )

    add_evidence_block(doc, [
        "\U0001f4c1 Folder: TPM Kiro Projects > Cost Savings Agent"
    ])

    # === SAVE ===
    output_path = Path(__file__).parent / "Growth_Timeline_Preview.docx"
    doc.save(str(output_path))
    print(f"Preview saved to: {output_path}")


if __name__ == "__main__":
    build_preview()
