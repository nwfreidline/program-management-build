"""
Build the One Sheet Document Template for DocForge.

Same formatting as the Narrative Template:
- Calibri fonts, narrow margins, compact table styling
- Header/footer with "Amazon Confidential" and page numbers
- Heading sizes: H1=16pt, H2=13pt, H3=11pt bold
- Body: 11pt, single-spaced, 6pt space after
- Tables: 10pt, light gray borders, shaded header row
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_style(doc, style_name, font_name="Calibri", font_size=11, bold=False,
              space_before=0, space_after=0, line_spacing=1.0,
              keep_with_next=False, color=None, italic=False):
    """Configure a paragraph style."""
    style = doc.styles[style_name]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = RGBColor(*color)

    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next

    return style


def create_style(doc, style_name, base_style="Normal", font_name="Calibri",
                 font_size=11, bold=False, space_before=0, space_after=0,
                 line_spacing=1.0, keep_with_next=False, color=None, italic=False):
    """Create a new paragraph style if it doesn't exist."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base_style]

    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = RGBColor(*color)

    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next

    return style


def add_page_number(paragraph):
    """Add 'Page X of Y' field codes to a paragraph."""
    run1 = paragraph.add_run("Page ")

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run2 = paragraph.add_run()
    run2._r.append(fldChar1)

    instrText1 = OxmlElement('w:instrText')
    instrText1.set(qn('xml:space'), 'preserve')
    instrText1.text = ' PAGE '
    run3 = paragraph.add_run()
    run3._r.append(instrText1)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run4 = paragraph.add_run()
    run4._r.append(fldChar2)

    run5 = paragraph.add_run(" of ")

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run6 = paragraph.add_run()
    run6._r.append(fldChar3)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = ' NUMPAGES '
    run7 = paragraph.add_run()
    run7._r.append(instrText2)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run8 = paragraph.add_run()
    run8._r.append(fldChar4)


def build_template():
    doc = Document()

    # === PAGE SETUP ===
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.54)
    section.right_margin = Inches(0.54)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)

    # === STYLES ===

    # Normal / Body text
    set_style(doc, "Normal",
              font_name="Calibri", font_size=11, bold=False,
              space_before=0, space_after=6, line_spacing=1.0)

    # Heading 1 — Main title / section headers
    set_style(doc, "Heading 1",
              font_name="Calibri", font_size=16, bold=True,
              space_before=12, space_after=0, line_spacing=1.0,
              keep_with_next=True, color=(0, 0, 0))

    # Heading 2 — Subsection headers
    set_style(doc, "Heading 2",
              font_name="Calibri", font_size=13, bold=True,
              space_before=12, space_after=0, line_spacing=1.0,
              keep_with_next=True, color=(0, 0, 0))

    # Heading 3 — Sub-subsection headers
    set_style(doc, "Heading 3",
              font_name="Calibri", font_size=11, bold=True,
              space_before=6, space_after=0, line_spacing=1.0,
              keep_with_next=True, color=(0, 0, 0))

    # Heading 4 — Minor headers
    set_style(doc, "Heading 4",
              font_name="Calibri", font_size=11, bold=True, italic=True,
              space_before=6, space_after=0, line_spacing=1.0,
              keep_with_next=True, color=(0, 0, 0))

    # List Paragraph — Bullet points
    list_style = set_style(doc, "List Paragraph",
                           font_name="Calibri", font_size=11, bold=False,
                           space_before=0, space_after=2, line_spacing=1.0)
    list_style.paragraph_format.left_indent = Inches(0.25)

    # AMZN H2 — Amazon-style heading (compatibility)
    create_style(doc, "AMZN H2", base_style="Heading 2",
                 font_name="Calibri", font_size=13, bold=True,
                 space_before=12, space_after=3, line_spacing=1.0,
                 keep_with_next=True, color=(0, 0, 0))

    # === HEADER ===
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    header_para = header.paragraphs[0]
    header_para.text = "DocForge \u2014 One Sheet"
    run = header_para.runs[0] if header_para.runs else header_para.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # === FOOTER ===
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    footer_para = footer.paragraphs[0]

    # Tab stop for right-aligned page number
    pf = footer_para.paragraph_format
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Inches(7.42), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # "Amazon Confidential" on left
    run_left = footer_para.add_run("Amazon Confidential")
    run_left.font.name = "Calibri"
    run_left.font.size = Pt(9)
    run_left.font.color.rgb = RGBColor(128, 128, 128)

    # Tab to right side
    footer_para.add_run("\t")

    # Page number on right
    add_page_number(footer_para)
    for run in footer_para.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # === PLACEHOLDER ===
    doc.add_paragraph("", style="Normal")

    # === SAVE ===
    output_path = "templates/docx/One_Sheet_Template.docx"
    doc.save(output_path)
    print(f"Template saved to: {output_path}")


if __name__ == "__main__":
    build_template()
