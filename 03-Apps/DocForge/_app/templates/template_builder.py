"""
Template Builder — Programmatically creates .docx templates from user-selected parameters.

Used by the DocForge "Create New" template creator window.
Can also be run standalone for testing:
    python template_builder.py --name "My Template" --font "Calibri" --size 11
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Configuration defaults
# ---------------------------------------------------------------------------

TEMPLATES_DIR = Path(__file__).parent / "docx"

PAGE_SIZES = {
    "Letter": (Inches(8.5), Inches(11)),
    "A4": (Inches(8.27), Inches(11.69)),
}

MARGINS = {
    "Narrow": (0.5, 0.5, 0.5, 0.5),
    "Normal": (1.0, 1.0, 1.0, 1.0),
    "Wide": (1.25, 1.25, 1.0, 1.0),
}

LINE_SPACING_MAP = {
    "Single": 1.0,
    "1.15": 1.15,
    "1.5": 1.5,
    "Double": 2.0,
}

HEADING_STYLES = {
    "Bold only": {"bold": True, "color": None, "underline": False},
    "Bold + color": {"bold": True, "color": (0, 51, 102), "underline": False},
    "Underlined": {"bold": True, "color": None, "underline": True},
}

FONTS = ["Calibri", "Arial", "Aptos", "Times New Roman", "Segoe UI", "Verdana"]
BODY_SIZES = [9, 10, 11, 12, 13, 14]


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _configure_style(style, font_name, font_size, bold=False, italic=False,
                     color=None, underline=False, space_before=0, space_after=6,
                     line_spacing=1.0, keep_with_next=False, left_indent=None):
    """Apply formatting to a paragraph style."""
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    font.underline = underline
    if color:
        font.color.rgb = RGBColor(*color)
    else:
        font.color.rgb = None

    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)


def _add_page_number(paragraph):
    """Add 'Page X of Y' field codes."""
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # PAGE field
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r = paragraph.add_run()
    r._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2 = paragraph.add_run()
    r2._r.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r3 = paragraph.add_run()
    r3._r.append(fld_end)

    run2 = paragraph.add_run(" of ")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    # NUMPAGES field
    fld_begin2 = OxmlElement('w:fldChar')
    fld_begin2.set(qn('w:fldCharType'), 'begin')
    r4 = paragraph.add_run()
    r4._r.append(fld_begin2)

    instr2 = OxmlElement('w:instrText')
    instr2.set(qn('xml:space'), 'preserve')
    instr2.text = ' NUMPAGES '
    r5 = paragraph.add_run()
    r5._r.append(instr2)

    fld_end2 = OxmlElement('w:fldChar')
    fld_end2.set(qn('w:fldCharType'), 'end')
    r6 = paragraph.add_run()
    r6._r.append(fld_end2)


# ---------------------------------------------------------------------------
# Reference doc extraction
# ---------------------------------------------------------------------------

def extract_styles_from_reference(ref_path: str) -> dict:
    """Extract key style properties from a reference .docx file.

    Returns a dict of parameters that can be passed to build_template().
    """
    doc = Document(ref_path)
    params = {}

    # Page setup from first section
    section = doc.sections[0]
    params["page_width"] = section.page_width
    params["page_height"] = section.page_height
    params["margin_left"] = section.left_margin
    params["margin_right"] = section.right_margin
    params["margin_top"] = section.top_margin
    params["margin_bottom"] = section.bottom_margin

    # Normal style (body text)
    try:
        normal = doc.styles["Normal"]
        if normal.font.name:
            params["font_name"] = normal.font.name
        if normal.font.size:
            params["font_size"] = int(normal.font.size.pt)
        if normal.paragraph_format.line_spacing:
            params["line_spacing"] = normal.paragraph_format.line_spacing
    except (KeyError, AttributeError):
        pass

    # Heading 1 style
    try:
        h1 = doc.styles["Heading 1"]
        if h1.font.bold:
            params["heading_bold"] = True
        if h1.font.color and h1.font.color.rgb:
            rgb = h1.font.color.rgb
            params["heading_color"] = (rgb[0], rgb[1], rgb[2]) if rgb else None
        if h1.font.size:
            params["heading1_size"] = int(h1.font.size.pt)
    except (KeyError, AttributeError):
        pass

    # Header/footer text
    try:
        header = section.header
        if header.paragraphs and header.paragraphs[0].text:
            params["header_text"] = header.paragraphs[0].text
    except (AttributeError, IndexError):
        pass

    try:
        footer = section.footer
        if footer.paragraphs and footer.paragraphs[0].text:
            footer_text = footer.paragraphs[0].text.strip()
            # Filter out page number fields (they show as empty or "PAGE")
            if footer_text and "PAGE" not in footer_text.upper():
                params["footer_text"] = footer_text
    except (AttributeError, IndexError):
        pass

    return params


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_template(
    name: str,
    font_name: str = "Calibri",
    body_size: int = 11,
    page_size: str = "Letter",
    margins: str = "Normal",
    line_spacing: str = "Single",
    heading_style: str = "Bold only",
    header_text: str = "",
    footer_text: str = "",
    show_page_numbers: bool = True,
    confidentiality: str = "None",
    reference_params: dict = None,
) -> str:
    """Build a .docx template from parameters.

    Args:
        name: Template name (used as filename)
        font_name: Body font family
        body_size: Body font size in pt
        page_size: "Letter" or "A4"
        margins: "Narrow", "Normal", or "Wide"
        line_spacing: "Single", "1.15", "1.5", or "Double"
        heading_style: "Bold only", "Bold + color", or "Underlined"
        header_text: Text for the page header (empty = no header)
        footer_text: Text for the left side of the footer
        show_page_numbers: Whether to add "Page X of Y" to footer
        confidentiality: "None", "Amazon Confidential", "Internal Only", or custom text
        reference_params: Dict from extract_styles_from_reference() to override defaults

    Returns:
        Path to the saved template file.
    """
    doc = Document()

    # Apply reference doc overrides if provided
    ref = reference_params or {}

    # === PAGE SETUP ===
    section = doc.sections[0]

    if "page_width" in ref:
        section.page_width = ref["page_width"]
        section.page_height = ref["page_height"]
    else:
        w, h = PAGE_SIZES.get(page_size, PAGE_SIZES["Letter"])
        section.page_width = w
        section.page_height = h

    if "margin_left" in ref:
        section.left_margin = ref["margin_left"]
        section.right_margin = ref["margin_right"]
        section.top_margin = ref["margin_top"]
        section.bottom_margin = ref["margin_bottom"]
    else:
        l, r, t, b = MARGINS.get(margins, MARGINS["Normal"])
        section.left_margin = Inches(l)
        section.right_margin = Inches(r)
        section.top_margin = Inches(t)
        section.bottom_margin = Inches(b)

    # Resolve parameters (reference overrides user selections)
    _font = ref.get("font_name", font_name)
    _size = ref.get("font_size", body_size)
    _spacing = ref.get("line_spacing", LINE_SPACING_MAP.get(line_spacing, 1.0))
    if isinstance(_spacing, str):
        _spacing = LINE_SPACING_MAP.get(_spacing, 1.0)

    # Heading config
    h_config = HEADING_STYLES.get(heading_style, HEADING_STYLES["Bold only"])
    if "heading_color" in ref:
        h_config = {**h_config, "color": ref["heading_color"]}
    h1_size = ref.get("heading1_size", _size + 5)

    # === STYLES ===

    # Normal
    _configure_style(
        doc.styles["Normal"],
        font_name=_font, font_size=_size,
        space_after=6, line_spacing=_spacing,
    )

    # Heading 1
    _configure_style(
        doc.styles["Heading 1"],
        font_name=_font, font_size=h1_size,
        bold=h_config["bold"], color=h_config["color"],
        underline=h_config["underline"],
        space_before=12, space_after=3, line_spacing=1.0,
        keep_with_next=True,
    )

    # Heading 2
    _configure_style(
        doc.styles["Heading 2"],
        font_name=_font, font_size=h1_size - 3,
        bold=h_config["bold"], color=h_config["color"],
        underline=h_config["underline"],
        space_before=10, space_after=2, line_spacing=1.0,
        keep_with_next=True,
    )

    # Heading 3
    _configure_style(
        doc.styles["Heading 3"],
        font_name=_font, font_size=_size,
        bold=True, color=h_config["color"],
        space_before=6, space_after=2, line_spacing=1.0,
        keep_with_next=True,
    )

    # List Paragraph
    _configure_style(
        doc.styles["List Paragraph"],
        font_name=_font, font_size=_size,
        space_after=2, line_spacing=_spacing,
        left_indent=0.25,
    )

    # === HEADER ===
    _header_text = ref.get("header_text", header_text)
    if _header_text:
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        hp = header.paragraphs[0]
        run = hp.add_run(_header_text)
        run.font.name = _font
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # === FOOTER ===
    _footer_text = ref.get("footer_text", footer_text)
    # Resolve confidentiality label
    if confidentiality and confidentiality != "None":
        _footer_text = confidentiality

    if _footer_text or show_page_numbers:
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        fp = footer.paragraphs[0]

        # Tab stop for right-aligned page number
        pf = fp.paragraph_format
        page_width_inches = section.page_width / Inches(1)
        margin_total = (section.left_margin + section.right_margin) / Inches(1)
        tab_pos = page_width_inches - margin_total
        pf.tab_stops.add_tab_stop(Inches(tab_pos), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

        if _footer_text:
            run = fp.add_run(_footer_text)
            run.font.name = _font
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

        if show_page_numbers:
            fp.add_run("\t")
            _add_page_number(fp)

    # === PLACEHOLDER (empty body so template isn't blank) ===
    doc.add_paragraph("", style="Normal")

    # === SAVE ===
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(c for c in name if c.isalnum() or c in " _-").strip()
    filename = f"{safe_name.replace(' ', '_')}_Template.docx"
    output_path = TEMPLATES_DIR / filename

    doc.save(str(output_path))
    return str(output_path)


# ---------------------------------------------------------------------------
# CLI for testing
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Build a DocForge template")
    parser.add_argument("--name", default="Test Template")
    parser.add_argument("--font", default="Calibri")
    parser.add_argument("--size", type=int, default=11)
    parser.add_argument("--page", default="Letter", choices=list(PAGE_SIZES.keys()))
    parser.add_argument("--margins", default="Normal", choices=list(MARGINS.keys()))
    parser.add_argument("--spacing", default="Single", choices=list(LINE_SPACING_MAP.keys()))
    parser.add_argument("--heading", default="Bold only", choices=list(HEADING_STYLES.keys()))
    parser.add_argument("--header", default="")
    parser.add_argument("--footer", default="")
    parser.add_argument("--confidentiality", default="None")
    parser.add_argument("--reference", default=None, help="Path to reference .docx")
    args = parser.parse_args()

    ref_params = None
    if args.reference:
        ref_params = extract_styles_from_reference(args.reference)
        print(f"Extracted styles from reference: {ref_params}")

    path = build_template(
        name=args.name,
        font_name=args.font,
        body_size=args.size,
        page_size=args.page,
        margins=args.margins,
        line_spacing=args.spacing,
        heading_style=args.heading,
        header_text=args.header,
        footer_text=args.footer,
        confidentiality=args.confidentiality,
        reference_params=ref_params,
    )
    print(f"Template created: {path}")
