"""
Build the Growth Timeline Template for DocForge.

Designed for converting growth_timeline_YYYY.md files into reader-friendly Word docs
suitable for performance reviews, promotion discussions, and compensation conversations.

Key formatting characteristics:
- Page: Letter size, narrow margins (0.54" L/R, ~0.72" T/B)
- Header: "DocForge — Growth Timeline" left-aligned, gray 9pt
- Footer: "Amazon Confidential" left + "Page X of Y" right
- Title: Calibri 18pt bold, centered — "Growth Timeline — [Year]"
- Subtitle: Calibri 11pt italic, centered — purpose statement
- Month headers: Calibri 16pt bold, dark blue (#003366), top border accent
- Project entries: Calibri 13pt bold, with status tag (italic, colored)
- STAR labels: Calibri 11pt bold inline, body text follows
- Action items: Calibri 11pt bullet list, 0.25" indent
- Supporting Evidence: Calibri 10pt italic, gray (#555555)
- Blockquotes: Calibri 10pt italic, 0.3" left indent, gray left border
- Section dividers: Thin horizontal rule between months
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_style(doc, style_name, font_name="Calibri", font_size=11, bold=False,
              space_before=0, space_after=0, line_spacing=1.0,
              keep_with_next=False, color=None, italic=False):
    """Configure a paragraph style."""
    style = doc.styles[style_name]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = RGBColor(*color)

    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next

    return style


def create_style(doc, style_name, base_style="Normal", font_name="Calibri",
                 font_size=11, bold=False, space_before=0, space_after=0,
                 line_spacing=1.0, keep_with_next=False, color=None, italic=False,
                 left_indent=None):
    """Create a new paragraph style if it doesn't exist."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base_style]

    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = RGBColor(*color)

    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)

    return style


def add_page_number(paragraph):
    """Add 'Page X of Y' field codes to a paragraph."""
    paragraph.add_run("Page ")

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run2 = paragraph.add_run()
    run2._r.append(fldChar1)

    instrText1 = OxmlElement('w:instrText')
    instrText1.set(qn('xml:space'), 'preserve')
    instrText1.text = ' PAGE '
    run3 = paragraph.add_run()
    run3._r.append(instrText1)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run4 = paragraph.add_run()
    run4._r.append(fldChar2)

    paragraph.add_run(" of ")

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run6 = paragraph.add_run()
    run6._r.append(fldChar3)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = ' NUMPAGES '
    run7 = paragraph.add_run()
    run7._r.append(instrText2)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run8 = paragraph.add_run()
    run8._r.append(fldChar4)


def add_bottom_border(paragraph, color="003366", size="8"):
    """Add a bottom border to a paragraph (used for month section dividers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_template():
    doc = Document()

    # === PAGE SETUP ===
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.54)
    section.right_margin = Inches(0.54)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)

    # === STYLES ===

    # Normal / Body text
    set_style(doc, "Normal",
              font_name="Calibri", font_size=11, bold=False,
              space_before=0, space_after=6, line_spacing=1.0)

    # Heading 1 — Section headers (e.g., "Project History", "Completed Projects")
    # Dark blue with bottom border accent
    set_style(doc, "Heading 1",
              font_name="Calibri", font_size=16, bold=True,
              space_before=18, space_after=6, line_spacing=1.0,
              keep_with_next=True, color=(0, 51, 102))

    # Heading 2 — Project/achievement entry titles
    set_style(doc, "Heading 2",
              font_name="Calibri", font_size=13, bold=True,
              space_before=12, space_after=2, line_spacing=1.0,
              keep_with_next=True, color=(0, 51, 102))

    # Heading 3 — Sub-sections within entries (e.g., "Supporting Evidence")
    set_style(doc, "Heading 3",
              font_name="Calibri", font_size=11, bold=True,
              space_before=6, space_after=2, line_spacing=1.0,
              keep_with_next=True, color=(0, 0, 0))

    # Heading 4 — Minor headers
    set_style(doc, "Heading 4",
              font_name="Calibri", font_size=11, bold=True, italic=True,
              space_before=6, space_after=0, line_spacing=1.0,
              keep_with_next=True, color=(0, 0, 0))

    # List Paragraph — Bullet points (action items, results)
    list_style = set_style(doc, "List Paragraph",
                           font_name="Calibri", font_size=11, bold=False,
                           space_before=0, space_after=2, line_spacing=1.0)
    list_style.paragraph_format.left_indent = Inches(0.5)

    # --- Custom styles for Growth Timeline ---

    # Title style — Document title (centered, large)
    create_style(doc, "GT Title", base_style="Normal",
                 font_name="Calibri", font_size=18, bold=True,
                 space_before=0, space_after=4, line_spacing=1.0,
                 color=(0, 51, 102))

    # Subtitle — Purpose statement under title
    create_style(doc, "GT Subtitle", base_style="Normal",
                 font_name="Calibri", font_size=11, italic=True,
                 space_before=0, space_after=12, line_spacing=1.0,
                 color=(85, 85, 85))

    # Status Tag — Inline status indicator (Complete, In Progress, etc.)
    create_style(doc, "GT Status", base_style="Normal",
                 font_name="Calibri", font_size=10, italic=True,
                 space_before=0, space_after=6, line_spacing=1.0,
                 color=(85, 85, 85))

    # STAR Label body — For STAR format content
    create_style(doc, "GT STAR", base_style="Normal",
                 font_name="Calibri", font_size=11, bold=False,
                 space_before=2, space_after=4, line_spacing=1.0,
                 left_indent=0.15)

    # Evidence / Blockquote — Supporting evidence callout
    create_style(doc, "GT Evidence", base_style="Normal",
                 font_name="Calibri", font_size=10, italic=True,
                 space_before=6, space_after=6, line_spacing=1.0,
                 color=(85, 85, 85), left_indent=0.3)

    # Section Summary — For the COMPLETED / BRAINSTORM lists
    create_style(doc, "GT Summary List", base_style="Normal",
                 font_name="Calibri", font_size=11, bold=False,
                 space_before=0, space_after=2, line_spacing=1.0,
                 left_indent=0.25)

    # === HEADER (none — clean top) ===
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    # === FOOTER ===
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    footer_para = footer.paragraphs[0]

    # Tab stop for right-aligned page number
    pf = footer_para.paragraph_format
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Inches(7.42), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # "Amazon Confidential" on left
    run_left = footer_para.add_run("Amazon Confidential")
    run_left.font.name = "Calibri"
    run_left.font.size = Pt(9)
    run_left.font.color.rgb = RGBColor(128, 128, 128)

    # Tab to right side
    footer_para.add_run("\t")

    # Page number on right
    add_page_number(footer_para)
    for run in footer_para.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # === PLACEHOLDER CONTENT ===
    # Empty paragraph so template isn't blank
    doc.add_paragraph("", style="Normal")

    # === SAVE ===
    from pathlib import Path
    script_dir = Path(__file__).parent
    docx_dir = script_dir / "docx"
    docx_dir.mkdir(parents=True, exist_ok=True)

    output_path = docx_dir / "Growth_Timeline_Template.docx"
    doc.save(str(output_path))
    print(f"Template saved to: {output_path}")

    # Also save at root templates dir
    root_copy = script_dir / "Growth_Timeline_Template.docx"
    doc.save(str(root_copy))
    print(f"Also saved to: {root_copy}")


if __name__ == "__main__":
    build_template()
